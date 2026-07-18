from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOFTWARE_FULL_NAME = "雀趣消除乐游戏软件"
SOFTWARE_SHORT_NAME = "雀趣消除乐"
VERSION = "V1.0"
COMPLETION_DATE = "2026年7月18日"
OUTPUT_NAMES = {
    "form": "01-软著申请信息采集表-雀趣消除乐V1.0.docx",
    "manual": "02-软件说明书-雀趣消除乐V1.0.docx",
    "source": "03-源程序鉴别材料-雀趣消除乐V1.0.docx",
    "checklist": "04-提交前核对清单-雀趣消除乐V1.0.docx",
}
SOURCE_FILES = [
    Path("assets/scripts/DuiDuiMahjongGame.ts"),
    Path("assets/scripts/model/DuiDuiMahjongModel.ts"),
    Path("assets/scripts/platform/DuiDuiAdConfig.ts"),
    Path("assets/scripts/platform/DuiDuiAdService.ts"),
    Path("assets/scripts/view/DuiDuiMahjongTheme.ts"),
]

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
MUTED = "666666"


def set_run_font(run, name="Microsoft YaHei", size=10.5, bold=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = sum(widths)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for old in list(grid):
        grid.remove(old)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def configure_document(doc: Document, title: str, subtitle: str = ""):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(f"{SOFTWARE_FULL_NAME} {VERSION}")
    set_run_font(run, size=9, color=MUTED)
    add_page_field(section.footer.paragraphs[0])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title)
    set_run_font(run, size=22, bold=True, color=DARK_BLUE)
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)
        run = p.add_run(subtitle)
        set_run_font(run, size=11, color=MUTED)


def add_info_table(doc, rows, widths=(2700, 6660), header=None):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    if header:
        cells = table.add_row().cells
        cells[0].merge(cells[1])
        cells[0].text = header
        set_cell_shading(cells[0], HEADER_FILL)
        for run in cells[0].paragraphs[0].runs:
            set_run_font(run, bold=True, color=DARK_BLUE)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], LIGHT_FILL)
        for run in cells[0].paragraphs[0].runs:
            set_run_font(run, bold=True)
        for run in cells[1].paragraphs[0].runs:
            set_run_font(run)
    set_table_widths(table, list(widths))
    doc.add_paragraph()
    return table


def add_bullet(doc, text, checked=False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(4)
    mark = "☐" if checked else "•"
    run = p.add_run(f"{mark} {text}")
    set_run_font(run)
    return p


def build_application_form(out_path: Path):
    doc = Document()
    configure_document(doc, "计算机软件著作权登记申请信息采集表", "个人申请｜提交前请补齐空白字段")
    add_info_table(doc, [
        ("软件全称", SOFTWARE_FULL_NAME),
        ("软件简称", SOFTWARE_SHORT_NAME),
        ("版本号", VERSION),
        ("软件分类", "应用软件｜游戏软件"),
        ("开发方式", "独立开发"),
        ("权利取得方式", "原始取得"),
        ("权利范围", "全部权利"),
        ("发表状态", "未发表"),
        ("开发完成日期", COMPLETION_DATE),
    ], header="一、软件基本信息")
    add_info_table(doc, [
        ("姓名", "________________________________"),
        ("身份证号码", "________________________________"),
        ("手机号", "________________________________"),
        ("电子邮箱", "________________________________"),
        ("通讯地址", "________________________________"),
        ("邮政编码", "________________________________"),
    ], header="二、著作权人信息（由申请人填写）")
    add_info_table(doc, [
        ("开发硬件环境", "普通个人计算机；请按实际设备补充 CPU、内存"),
        ("运行硬件环境", "支持抖音客户端的智能手机"),
        ("开发操作系统", "Windows；请按实际版本补充"),
        ("开发工具", "Cocos Creator 3.8.7、TypeScript"),
        ("运行平台", "抖音小游戏运行环境"),
        ("编程语言", "TypeScript"),
        ("源程序量", "以最终提交的源程序鉴别材料统计为准"),
        ("开发目的", "提供轻松易上手的麻将主题消除游戏体验"),
        ("面向领域", "休闲益智游戏"),
        ("主要功能", "麻将牌堆消除、提示、撤回、洗牌、音效与设置、广告激励道具"),
    ], header="三、软件技术信息")
    doc.add_heading("四、申请前确认", level=1)
    for item in [
        "上述软件名称、简称和版本号与申请系统填写内容完全一致。",
        "个人身份信息与身份证件完全一致，证件在有效期内。",
        "提交的说明书和源程序材料均来自本人拥有著作权的软件。",
        "游戏使用的字体、音乐、美术和其他素材具有合法来源或授权。",
    ]:
        add_bullet(doc, item, checked=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.add_run("申请人签字：____________________    日期：______年____月____日")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def build_user_manual(out_path: Path):
    doc = Document()
    configure_document(doc, "软件说明书", f"{SOFTWARE_FULL_NAME} {VERSION}")
    doc.add_heading("1. 软件概述", level=1)
    doc.add_paragraph(
        "雀趣消除乐是一款麻将主题的休闲益智消除小游戏。玩家从层叠牌堆中选择未被遮挡的麻将牌，"
        "将相同牌型组合并消除，在有限槽位和牌堆条件下完成关卡目标。软件强调直观操作、轻量对局和即时反馈。"
    )
    add_info_table(doc, [
        ("软件名称", SOFTWARE_FULL_NAME),
        ("软件简称", SOFTWARE_SHORT_NAME),
        ("版本", VERSION),
        ("开发工具", "Cocos Creator 3.8.7"),
        ("主要语言", "TypeScript"),
        ("目标平台", "抖音小游戏"),
    ])
    doc.add_heading("2. 开发目的", level=1)
    doc.add_paragraph(
        "本软件以大众熟悉的麻将牌图形为主题，通过牌堆遮挡、同类匹配和道具辅助构成轻量策略玩法，"
        "让用户能够利用碎片时间获得清晰、轻松且具有阶段目标的娱乐体验。"
    )
    doc.add_heading("3. 运行环境", level=1)
    for item in [
        "客户端：安装并可正常运行抖音客户端的智能手机。",
        "网络：基础游戏逻辑可在本地运行；广告激励等平台能力需要网络连接。",
        "显示与输入：竖屏触控操作。",
        "开发环境：Windows、Cocos Creator 3.8.7、TypeScript。",
    ]:
        add_bullet(doc, item)
    doc.add_heading("4. 主要功能", level=1)
    features = [
        ("首页与模式入口", "展示游戏名称、主要入口和设置入口，用户可进入消除对局。"),
        ("麻将牌堆生成", "根据关卡数据生成具有层级和遮挡关系的麻将牌堆。"),
        ("可选牌判断", "只有未被上层牌遮挡且可操作的麻将牌能够被选中。"),
        ("槽位与匹配消除", "选中的牌进入槽位；相同牌型达到匹配条件后触发消除与反馈。"),
        ("胜负与重开", "牌堆清空时完成关卡；槽位达到限制且无法消除时结束本局，可重新开始。"),
        ("提示道具", "高亮当前局面中可形成有效操作的麻将牌。"),
        ("撤回道具", "在规则允许时撤销最近一次选择并恢复对应状态。"),
        ("洗牌道具", "重新排列剩余可用麻将牌，帮助用户调整局面。"),
        ("设置与音频", "提供背景音乐、音效等开关，并保存用户设置。"),
        ("广告激励", "在用户确认后调用平台激励视频能力，成功观看后发放相应道具。"),
    ]
    add_info_table(doc, features, widths=(2400, 6960))
    doc.add_heading("5. 操作流程", level=1)
    steps = [
        "启动软件并进入首页。",
        "点击开始入口进入消除关卡。",
        "观察牌堆遮挡关系，点击当前可选的麻将牌。",
        "选中牌进入下方槽位；满足同类匹配条件时自动消除。",
        "需要时使用提示、撤回或洗牌道具；涉及激励视频时先确认观看。",
        "清空牌堆后完成关卡；槽位达到限制时可重新开始。",
        "通过设置入口调整背景音乐与音效。",
    ]
    for idx, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.add_run(f"{idx}. {step}")
    doc.add_heading("6. 核心规则说明", level=1)
    doc.add_paragraph(
        "软件通过麻将牌的空间层级计算遮挡关系。用户只能选择当前可操作的牌。选中牌按照规则进入槽位，"
        "系统持续检查同类牌数量并执行消除；同时维护剩余牌、槽位状态、道具次数和对局结果。"
    )
    doc.add_heading("7. 数据与异常处理", level=1)
    for item in [
        "用户设置保存于本地存储，重新进入软件时恢复。",
        "广告未加载、播放失败或用户中途退出时，不发放对应激励并给出状态提示。",
        "弹窗和遮罩用于阻止重复输入，避免同一操作被连续触发。",
        "重新开始时清理本局临时状态并重新生成牌堆。",
    ]:
        add_bullet(doc, item)
    doc.add_heading("8. 界面截图补充清单", level=1)
    p = doc.add_paragraph()
    set_cell_text = p.add_run("正式提交前请从实际运行版本截取以下画面并插入本节；不得使用概念图代替实际界面。")
    set_run_font(set_cell_text, bold=True, color="9B1C1C")
    for item in [
        "软件首页：完整显示“雀趣消除乐”名称和开始入口。",
        "对局主界面：显示麻将牌堆、槽位和道具按钮。",
        "匹配消除反馈：显示同类牌消除后的局面。",
        "设置界面：显示背景音乐、音效等开关。",
        "关卡完成或失败界面：显示结果与重新开始入口。",
    ]:
        add_bullet(doc, item, checked=True)
    doc.add_heading("9. 技术特点", level=1)
    for item in [
        "使用 TypeScript 组织游戏状态、牌堆模型、平台广告服务和界面主题。",
        "模型层维护牌型、层级、遮挡、槽位和匹配状态，界面层根据状态刷新显示。",
        "平台能力通过独立广告服务封装，降低核心玩法对具体平台接口的依赖。",
        "采用程序化界面节点与主题参数，便于适配竖屏小游戏场景。",
    ]:
        add_bullet(doc, item)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def collect_source_lines(project_root: Path):
    collected = []
    for relative in SOURCE_FILES:
        path = project_root / relative
        lines = path.read_text(encoding="utf-8-sig").splitlines()
        collected.append((relative.as_posix(), lines))
    return collected


def build_source_listing(out_path: Path, project_root: Path):
    sources = collect_source_lines(project_root)
    flat = []
    for filename, lines in sources:
        flat.append(f"// ===== 文件：{filename} =====")
        flat.extend(line.rstrip() for line in lines)
        flat.append("")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(f"{SOFTWARE_FULL_NAME} {VERSION}｜源程序鉴别材料")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_field(section.footer.paragraphs[0])

    lines_per_page = 50
    for page_start in range(0, len(flat), lines_per_page):
        page_lines = flat[page_start:page_start + lines_per_page]
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        for index, line in enumerate(page_lines):
            absolute = page_start + index + 1
            run = p.add_run(f"{absolute:04d}  {line}")
            set_run_font(run, name="Consolas", size=7.5)
            if index < len(page_lines) - 1:
                run.add_break()
        if page_start + lines_per_page < len(flat):
            doc.add_page_break()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return len(flat), (len(flat) + lines_per_page - 1) // lines_per_page


def build_submission_checklist(out_path: Path):
    doc = Document()
    configure_document(doc, "软著申请与抖音上架提交前核对清单", f"{SOFTWARE_FULL_NAME} {VERSION}")
    sections = [
        ("一、申请前", [
            "软件全称统一为“雀趣消除乐游戏软件”，简称统一为“雀趣消除乐”，版本号统一为 V1.0。",
            "完成日期填写 2026年7月18日，发表状态选择“未发表”。",
            "姓名、身份证号码、手机号、邮箱、地址与个人证件信息一致。",
            "确认字体、音乐、音效、美术、图标和其他素材具有原创证明或合法授权。",
        ]),
        ("二、软著提交", [
            "在线申请表、软件说明书和源程序鉴别材料中的名称、简称、版本号完全一致。",
            "源程序为自有代码，不含第三方库、密钥、账号、个人隐私或无关项目名称。",
            "源程序页眉标注软件名称和版本号，页码连续，代码清晰可读。",
            "软件说明书插入实际运行截图，截图内游戏名称已更新为“雀趣消除乐”。",
            "个人申请人按办理渠道要求完成实名认证、签章或身份核验。",
        ]),
        ("三、证书领取", [
            "核对证书上的软件名称、简称、版本号、登记号、登记日期和著作权人。",
            "保存证书原件、清晰扫描件和官方查询结果。",
            "不得修改证书内容，不添加无关水印。",
        ]),
        ("四、抖音资质准入", [
            "抖音小游戏名称使用“雀趣消除乐”，与软著简称完全一致。",
            "同一软著只用于一款在抖音平台在线运营的小游戏。",
            "若抖音平台主体不是软著登记的个人著作权人，准备完整、清晰且符合平台要求的运营/发行授权证明。",
            "上传软著证书清晰原件或符合要求的扫描件，确保可通过登记号核验。",
            "如游戏涉及第三方 IP、商标、音乐或美术作品，另行上传对应权属证明和授权链。",
        ]),
        ("五、小游戏备案与后续", [
            "备案前准备国家版权局发放的软著登记证书；不要仅依赖电子软著准入结果。",
            "备案、ICP备案、平台名称和游戏内展示名称保持一致。",
            "若后续更换软著或名称，先评估重新备案和资质修改次数影响。",
            "如开通内购能力，另行核对网络游戏出版物号（版号）及运营单位要求。",
        ]),
    ]
    for heading, items in sections:
        doc.add_heading(heading, level=1)
        for item in items:
            add_bullet(doc, item, checked=True)
    doc.add_heading("官方参考", level=1)
    for item in [
        "抖音小游戏基础信息审核规范：developer.open-douyin.com/docs/resource/zh-CN/mini-game/operation1/norms/game-info-audit",
        "抖音小游戏资质规范：developer.open-douyin.com/docs/resource/zh-CN/mini-game/operation1/norms/credential-norms-for-mini-game",
        "抖音小游戏审核常见问题（资质篇）：developer.open-douyin.com/docs/resource/zh-CN/mini-game/guide/contactus/endowments",
    ]:
        add_bullet(doc, item)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def validate_outputs(paths):
    expected_tokens = {
        OUTPUT_NAMES["form"]: ["计算机软件著作权登记申请信息采集表", SOFTWARE_FULL_NAME, "身份证号码"],
        OUTPUT_NAMES["manual"]: ["软件说明书", SOFTWARE_FULL_NAME, "主要功能"],
        OUTPUT_NAMES["source"]: ["DuiDuiMahjongGame", "雀趣消除乐"],
        OUTPUT_NAMES["checklist"]: ["核对清单", SOFTWARE_SHORT_NAME, "抖音资质准入"],
    }
    for path in paths:
        assert path.exists() and path.stat().st_size > 0, path
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            text += "\n" + "\n".join(cell.text for row in table.rows for cell in row.cells)
        for token in expected_tokens[path.name]:
            assert token in text, (path, token)
    form = Document(paths[0])
    form_text = "\n".join(cell.text for t in form.tables for row in t.rows for cell in row.cells)
    assert "________________________________" in form_text


def build_all(output_dir: Path, project_root: Path):
    output_dir.mkdir(parents=True, exist_ok=True)
    paths = [output_dir / OUTPUT_NAMES[key] for key in ("form", "manual", "source", "checklist")]
    build_application_form(paths[0])
    build_user_manual(paths[1])
    line_count, page_count = build_source_listing(paths[2], project_root)
    build_submission_checklist(paths[3])
    validate_outputs(paths)
    print(f"源程序：{line_count} 行，预计 {page_count} 页")
    for path in paths:
        print(path)
    return paths


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    project_root = Path(__file__).resolve().parents[1]
    output = args.output if args.output.is_absolute() else project_root / args.output
    build_all(output, project_root)


if __name__ == "__main__":
    main()
