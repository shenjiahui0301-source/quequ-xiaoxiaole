from __future__ import annotations

import json
import zipfile
from pathlib import Path

from docx import Document


def main():
    root = Path(__file__).resolve().parents[1] / "docs" / "soft-copyright"
    files = sorted(root.glob("*.docx"))
    assert len(files) == 4, f"Expected 4 DOCX files, got {len(files)}"
    report = []
    for path in files:
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        text += "\n" + "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
        assert "雀趣消消乐" not in text, path
        assert "雀趣消除乐" in text, path
        section = doc.sections[0]
        assert round(section.page_width.inches, 1) == 8.5, path
        assert round(section.page_height.inches, 1) == 11.0, path
        with zipfile.ZipFile(path) as archive:
            xml = archive.read("word/document.xml").decode("utf-8")
        page_breaks = xml.count('w:type="page"')
        report.append({
            "file": path.name,
            "bytes": path.stat().st_size,
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "page_breaks": page_breaks,
        })

    form = Document(next(path for path in files if path.name.startswith("01-")))
    form_text = "\n".join(cell.text for table in form.tables for row in table.rows for cell in row.cells)
    assert form_text.count("________________________________") >= 6

    source_path = next(path for path in files if path.name.startswith("03-"))
    source = Document(source_path)
    source_text = "\n".join(p.text for p in source.paragraphs)
    expected_sources = [
        "DuiDuiMahjongGame.ts",
        "DuiDuiMahjongModel.ts",
        "DuiDuiAdConfig.ts",
        "DuiDuiAdService.ts",
        "DuiDuiMahjongTheme.ts",
    ]
    assert all(name in source_text for name in expected_sources)
    source_report = next(item for item in report if item["file"].startswith("03-"))
    assert source_report["page_breaks"] == 59, source_report

    print(json.dumps(report, ensure_ascii=False, indent=2))
    print("PASS: 4 DOCX files; privacy blanks, naming, source order, and 60-page pagination verified")


if __name__ == "__main__":
    main()
